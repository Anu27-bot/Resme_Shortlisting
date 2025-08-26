from dotenv import load_dotenv
load_dotenv()
import os
import subprocess
import pandas as pd
import re
import time
import logging
import io
import json
import hashlib
from datetime import datetime, timedelta
from dateutil.parser import parse
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from base64 import urlsafe_b64decode
from tabulate import tabulate
from docx import Document
from PyPDF2 import PdfReader
import shutil
import tempfile
import pytz
from pdfminer.high_level import extract_text
from groq import Groq
from supabase import create_client, Client
from fuzzywuzzy import fuzz
import logging
import numpy as np

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Configuration variables
script_dir = os.getcwd()
RESUME_FOLDER = os.path.join(script_dir, "Resumes")
LIBREOFFICE_PATH = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
OUTPUT_CSV = os.path.join(RESUME_FOLDER, "resume_analysis.csv")

# Groq Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable not set. Please set it in your .env file.")
client = Groq(api_key=GROQ_API_KEY)
GROQ_MODEL = 'llama3-8b-8192'

CURRENT_DATE = datetime.now().strftime("%Y-%m-%d")

# Supabase Configuration
SUPABASE_URL = "https://hpbopbkscnafszmrggcb.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImhwYm9wYmtzY25hZnN6bXJnZ2NiIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NTM5MDk0OTIsImV4cCI6MjA2OTQ4NTQ5Mn0.yeUpZUz2bWrjPKjTg8Qzxg2S2FbAMXeU10r5_i-vei4"

def remove_duplicate_candidates(df):
    """Remove duplicate candidates based on name, location, and experience."""
    # Create a hash for each candidate based on key fields
    df['CandidateHash'] = df.apply(lambda row: hashlib.md5(
        f"{row['Name']}{row['Current Location']}{row['Experience']}".encode()
    ).hexdigest(), axis=1)
    
    # Keep the first occurrence of each duplicate
    df = df.drop_duplicates(subset=['CandidateHash'], keep='first')
    
    # Recalculate ranks after removing duplicates
    df = df.sort_values(by=['Rank', 'Matching Skills Count'], ascending=[True, False])
    df['Rank'] = df['Composite Score'].rank(ascending=False, method='min').astype(int)
    
    return df.drop(columns=['CandidateHash'])

def auto_authenticate_google():
    """Authenticate and build Gmail API service."""
    creds = None
    token_path = 'token.json'
    SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
    if os.path.exists(token_path):
        try:
            creds = Credentials.from_authorized_user_file(token_path, SCOPES)
            if creds.expired and creds.refresh_token:
                creds.refresh(Request())
        except Exception as e:
            logging.error(f"Error loading credentials: {e}")
            creds = None
    if not creds or not creds.valid:
        try:
            flow = InstalledAppFlow.from_client_secrets_file('client.json', SCOPES)
            creds = flow.run_local_server(port=8080)
            token_data = json.loads(creds.to_json())
            token_data['creation_time'] = datetime.now(pytz.UTC).isoformat()
            with open(token_path, 'w') as token:
                json.dump(token_data, token)
        except Exception as e:
            logging.error(f"Authentication failed: {e}")
            raise
    try:
        gmail_service = build('gmail', 'v1', credentials=creds)
        return gmail_service
    except Exception as e:
        logging.error(f"Failed to build services: {e}")
        raise

def get_emails_by_job_id(service, job_id):
    """Fetch all emails containing the given job ID within the last 30 days."""
    today = datetime.now()
    last_month = today - timedelta(days=30)
    last_month_str = last_month.strftime('%Y/%m/%d')
    all_messages = []
    page_token = None
    
    while True:
        try:
            results = service.users().messages().list(
                userId="me",
                q=f"after:{last_month_str} {job_id}",
                maxResults=500,
                pageToken=page_token
            ).execute()
            messages = results.get("messages", [])
            all_messages.extend(messages)
            page_token = results.get("nextPageToken")
            if not page_token:
                break
            time.sleep(1)
        except Exception as e:
            logging.error(f"Error fetching emails: {e}")
            break
    return all_messages

def decode_base64(data):
    """Decode base64 email content safely."""
    missing_padding = len(data) % 4
    if missing_padding:
        data += '=' * (4 - missing_padding)
    return urlsafe_b64decode(data)

def extract_email_body(payload):
    """Extract email body content (text/plain or text/html)."""
    if not payload:
        return ""
    if "body" in payload and "data" in payload["body"]:
        return decode_base64(payload["body"]["data"]).decode("utf-8", errors="ignore")
    if "parts" in payload:
        for part in payload["parts"]:
            if part.get("mimeType", "") in ["text/plain", "text/html"] and "data" in part.get("body", {}):
                return decode_base64(part["body"]["data"]).decode("utf-8", errors="ignore")
            if "parts" in part:
                nested_body = extract_email_body(part)
                if nested_body:
                    return nested_body
    return ""

def clean_skill_text(skill_text):
    """Cleans formatting while retaining skill text with years."""
    if not skill_text:
        return None

    # Remove numbering/bullets while preserving other text
    skill_text = re.sub(r"(\n\s*[-•*]\s*\d+\.?\s*|\n\s*[-•*]\s*)", "\n", skill_text)
    skill_text = re.sub(r"^\d+\.\s*", "", skill_text, flags=re.MULTILINE)

    # Split into individual skills based on newlines or specific patterns
    skills = []
    current_skill = ""
    for line in skill_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Detect new skill (keywords like "Ability", "Understanding", "Strong", etc.)
        if re.match(r"^(Ability|Understanding|Strong|Experience|Knowledge|Excellent)", line, re.IGNORECASE):
            if current_skill:
                skills.append(current_skill.strip())
            current_skill = line
        else:
            current_skill += " " + line

    if current_skill:
        skills.append(current_skill.strip())

    return skills

def extract_skills(email_body):
    """Extracts Skills section while preserving 'Required X Years' text."""
    # First try to find a clear skills section
    match = re.search(
        r"Skills?:\s*(.*?)\s*(?:Responsibilities:|Qualifications:|Description:|Job ID:|$)",
        email_body,
        re.DOTALL | re.IGNORECASE
    )
    if match:
        skill_text = match.group(1).strip()
        return clean_skill_text(skill_text)
    
    # If no clear section, look for bulleted lists that might be skills
    bullet_items = re.findall(r"(?:\n\s*[-•*]\s*)(.+?)(?=\n\s*[-•*]|$)", email_body)
    if bullet_items:
        return clean_skill_text("\n".join(bullet_items))
    
    return None

def fetch_attachments(service, message_id):
    """Fetch all attachments metadata from an email."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg.get("payload", {})
        parts = payload.get("parts", [])
        attachments = []
        for part in parts:
            filename = part.get("filename")
            attachment_id = part.get("body", {}).get("attachmentId")
            if filename and attachment_id:
                attachments.append((filename, attachment_id))
        logging.info(f"Fetched {len(attachments)} attachments for message {message_id}")
        return attachments
    except Exception as e:
        logging.error(f"Error fetching attachments for message {message_id}: {e}")
        return []

def get_attachment_data(service, message_id, attachment_id):
    """Fetch and decode attachment data using Gmail API."""
    try:
        attachment = service.users().messages().attachments().get(
            userId="me",
            messageId=message_id,
            id=attachment_id
        ).execute()
        data = attachment.get("data", "")
        return urlsafe_b64decode(data)
    except Exception as e:
        logging.error(f"Error fetching attachment data for attachment {attachment_id}: {e}")
        return None

def extract_text_from_attachment(attachment_data, filename, temp_dir=None):
    """Extract text content from PDF, DOCX, or DOC attachments with fallback."""
    try:
        if filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(attachment_data))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        elif filename.lower().endswith(".pdf"):
            try:
                text = extract_text(io.BytesIO(attachment_data))
                if text.strip():
                    return text
            except Exception as e:
                logging.warning(f"pdfminer failed for {filename}: {e}")
            # Fallback to PyPDF2
            try:
                reader = PdfReader(io.BytesIO(attachment_data))
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
            except Exception as e:
                logging.warning(f"PyPDF2 failed for {filename}: {e}")
                return None
        elif filename.lower().endswith(".doc"):
            if not temp_dir:
                temp_dir = tempfile.mkdtemp()
            temp_doc_path = os.path.join(temp_dir, "temp.doc")
            with open(temp_doc_path, "wb") as f:
                f.write(attachment_data)
            for attempt in range(2):
                try:
                    docx_path = convert_doc_to_docx(temp_doc_path, temp_dir)
                    if docx_path.endswith(".docx"):
                        doc = Document(docx_path)
                        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
                        os.remove(docx_path)
                        os.remove(temp_doc_path)
                        return text
                    os.remove(temp_doc_path)
                    return None
                except Exception as e:
                    logging.warning(f"Attempt {attempt+1} failed for {filename}: {e}")
                    time.sleep(1)
            os.remove(temp_doc_path)
            return None
        return None
    except Exception as e:
        logging.error(f"Error extracting text from {filename}: {e}")
        return None
    finally:
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)

def is_potential_resume(filename):
    """Heuristic to check if filename might be a resume."""
    lower_name = filename.lower()
    EXCLUDED_KEYPHRASES = [
        "dl", "visa", "h1", "gc", "i-129", "approval", "sm", "skill matrix", "rtr", "innosoul",
        "reference", "patibandla", "check form", "sow", "ead", "70125071", "scanned",
        "driver license", "driving license", "passport", "i9", "w2", "paystub", "offer letter",
        "contract", "background check", "ssn", "social security", "id card", "certification form",
        "self-certification", "authorization form", "approval form", "clearance form",
        "verification form", "compliance form", "disclosure form", "attestation form",
        "acknowledgment form", "agreement form", "consent form", "declaration form",
        "enrollment form", "registration form", "application form", "submission form",
        "request form", "approval form", "clearance certificate", "security clearance",
        "background form", "screening form"
    ]
    if any(term in lower_name for term in EXCLUDED_KEYPHRASES):
        return False
    if not lower_name.endswith((".pdf", ".docx", ".doc")):
        return False
    return True

def is_resume_content(text):
    """Check if text content looks like a resume by keyword and section scoring."""
    if not text:
        return False
    NON_RESUME_PATTERNS = [
        r"reference\s*check", r"visa\s*status", r"IMG_\d{8}_\d{4}", r"approval\s*notice",
        r"skill\s*matrix", r"return\s*to\s*recruiter", r"form\s*[0-9]{3}",
        r"government\s*issued", r"validity\s*date", r"solicitation\s*number",
        r"candidate\s*reference"
    ]
    RESUME_SECTIONS = [
        r"work\s*experience", r"professional\s*(history|experience|summary)",
        r"skills?", r"education", r"projects?", r"certifications?",
        r"technical\s*(skills|proficiencies)", r"employment\s*history",
        r"key\s*skills", r"executive\s*summary", r"work\s*history",
        r"technical\s*summary", r"professional\s*overview"
    ]
    POSITIVE_KEYWORDS = [
        r"\d+\+?\s*years?\s*of\s*experience", r"developed", r"implemented",
        r"power\s*apps", r"power\s*automate", r"power\s*bi", r"azure",
        r"dataverse", r"microsoft\s*365", r"dynamics\s*365", r"sharepoint",
        r"certified", r"bachelor", r"master", r"engineer", r"developer",
        r"architect", r"solution", r"automation", r"integration"
    ]
    text = re.sub(r"\s+", " ", text.lower())
    if any(re.search(pattern, text) for pattern in NON_RESUME_PATTERNS):
        logging.debug("Non-resume content detected")
        return False
    section_count = sum(1 for pattern in RESUME_SECTIONS if re.search(pattern, text))
    keyword_count = sum(1 for pattern in POSITIVE_KEYWORDS if re.search(pattern, text))
    score = section_count * 2 + keyword_count
    return score >= 3

def validate_resume(service, message_id, attachment_id, filename):
    """Check if an attachment is likely a resume through filename and content."""
    if not is_potential_resume(filename):
        return False
    data = get_attachment_data(service, message_id, attachment_id)
    if not data:
        return False
    text = extract_text_from_attachment(data, filename)
    if not text:
        return False
    return is_resume_content(text)

def identify_resume(service, message_id, attachments):
    """
    Identify best resume attachment using filename heuristics and content validation.
    Returns filename or "N/A" if none found.
    """
    priority_candidates = []
    other_candidates = []

    for filename, attachment_id in attachments:
        lower_name = filename.lower()
        is_standard_resume = (
            "resume" in lower_name or 
            "cv" in lower_name or 
            "curriculum vitae" in lower_name or
            "bio data" in lower_name
        )
        
        name_pattern = (
            re.match(r"^[A-Z][a-z]+[A-Z][a-z]+\.(pdf|docx|doc)$", filename) or
            re.match(r"^[A-Z][a-z]+_[A-Z][a-z]+\.(pdf|docx|doc)$", filename) or
            "profile" in lower_name or
            "portfolio" in lower_name
        )
        valid_extension = lower_name.endswith((".pdf", ".docx", ".doc"))
        if (is_standard_resume or name_pattern) and valid_extension:
            priority_candidates.append((filename, attachment_id))
        elif valid_extension:
            other_candidates.append((filename, attachment_id))
        
    for filename, attachment_id in priority_candidates:
        if validate_resume(service, message_id, attachment_id, filename):
            logging.info(f"Identified resume by filename: {filename}")
            return filename

    validated_resumes = []
    for filename, attachment_id in other_candidates:
        if validate_resume(service, message_id, attachment_id, filename):
            validated_resumes.append(filename)

    if len(validated_resumes) == 1:
        logging.info(f"Identified resume by content: {validated_resumes[0]}")
        return validated_resumes[0]
    elif len(validated_resumes) > 1:
        best_candidate = None
        highest_score = 0
        for filename in validated_resumes:
            attachment_id = next(aid for (fn, aid) in other_candidates if fn == filename)
            data = get_attachment_data(service, message_id, attachment_id)
            text = extract_text_from_attachment(data, filename)
            if not text:
                continue
            score = 0
            resume_sections = [
                r"work\s*experience", r"professional\s*(history|experience|summary)",
                r"skills?", r"education", r"projects?", r"certifications?",
                r"technical\s*(skills|proficiencies)", r"employment\s*history",
                r"summary\s*of\s*qualifications", r"career\s*objective"
            ]
            for section in resume_sections:
                if re.search(section, text, re.IGNORECASE):
                    score += 1
            if score > highest_score:
                highest_score = score
                best_candidate = filename
        if best_candidate:
            logging.info(f"Selected best resume from multiple: {best_candidate}")
            return best_candidate
        return validated_resumes[0]

    if other_candidates:
        logging.warning(f"No clear resume found, returning first candidate: {other_candidates[0][0]}")
        return other_candidates[0][0]

    logging.warning("No valid resume found in attachments")
    return "N/A"

def save_resumes_to_folder(service, details, message_id, attachments, resume_folder):
    """Save identified resume attachment to folder."""
    resume_filename = details["Resume File"]
    if resume_filename == "N/A":
        logging.info(f"No resume file found for candidate: {details.get('Name', 'Unknown')}")
        return
    safe_filename = re.sub(r'[^\w\s.-]', '_', resume_filename)
    safe_filename = re.sub(r'\s+', '_', safe_filename)
    attachment_id = next((aid for (fn, aid) in attachments if fn == resume_filename), None)
    if not attachment_id:
        logging.warning(f"Attachment ID not found for resume file: {resume_filename}")
        return
    attachment_data = get_attachment_data(service, message_id, attachment_id)
    if not attachment_data:
        logging.error(f"Failed to fetch attachment data for resume file: {resume_filename}")
        return
    resume_file_path = os.path.join(resume_folder, safe_filename)
    try:
        with open(resume_file_path, "wb") as file:
            file.write(attachment_data)
        logging.info(f"Saved resume file: {resume_file_path}")
    except Exception as e:
        logging.error(f"Failed to save resume file {safe_filename}: {e}")

def create_resume_folder(base_folder="Resumes", job_id=None):
    """Create or clear folder for resumes with job ID subfolder."""
    if job_id:
        folder_name = os.path.join(base_folder, f"Job_{job_id}")
    else:
        folder_name = base_folder
        
    if os.path.exists(folder_name):
        for filename in os.listdir(folder_name):
            file_path = os.path.join(folder_name, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                logging.error(f"Failed to delete {file_path}: {e}")
    else:
        os.makedirs(folder_name, exist_ok=True)
    return folder_name

def convert_doc_to_docx(doc_path, output_folder):
    """Convert DOC to DOCX using LibreOffice headless."""
    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_folder,
            doc_path
        ], check=True, timeout=30)
        docx_file = os.path.join(output_folder, os.path.splitext(os.path.basename(doc_path))[0] + ".docx")
        if os.path.exists(docx_file):
            logging.info(f"Converted {doc_path} to {docx_file}")
            return docx_file
        logging.warning(f"DOCX conversion failed for {doc_path}: Output file not found")
        return doc_path
    except subprocess.CalledProcessError as e:
        logging.error(f"Error converting {doc_path} to DOCX: {e}")
        return doc_path
    except subprocess.TimeoutExpired:
        logging.error(f"Timeout converting {doc_path} to DOCX")
        return doc_path

def convert_docx_to_txt(input_path, output_path):
    """Convert DOCX to TXT with deep extraction including tables and headers."""
    try:
        doc = Document(input_path)
        full_text = []

        def extract_paragraphs(paragraphs):
            for para in paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

        extract_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                nested_text = nested_cell.text.strip()
                                if nested_text:
                                    full_text.append(nested_text)

        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    full_text.append(header.text.strip())
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    full_text.append(footer.text.strip())
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

        for shape in doc.element.body.iter():
            if shape.tag.endswith('wps:txbx'):
                for p in shape.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    text = ''.join(t.text for t in p.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) if t.text)
                    if text.strip():
                        full_text.append(text.strip())

        seen = set()
        full_text = [t for t in full_text if not (t in seen or seen.add(t))]

        if not full_text:
            logging.warning(f"No text extracted from {input_path}")
            full_text = ["No text content found in document"]

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))

        file_size = os.path.getsize(output_path)
        if file_size < 100:
            logging.warning(f"Generated {output_path} is very small ({file_size} bytes), may not contain full text")

        logging.info(f"Successfully converted {input_path} to {output_path} ({file_size} bytes)")
    except Exception as e:
        logging.error(f"Error converting {input_path} to .txt: {e}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Error extracting text: {str(e)}")

def convert_pdf_to_txt(input_path, output_path):
    """Convert PDF to TXT using pdfminer with PyPDF2 fallback."""
    try:
        text = extract_text(input_path)
        if text.strip():
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            logging.info(f"Successfully converted {input_path} to {output_path}")
            return
    except Exception as e:
        logging.warning(f"pdfminer failed for {input_path}: {e}")
    try:
        reader = PdfReader(input_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        logging.info(f"Successfully converted {input_path} to {output_path} using PyPDF2")
    except Exception as e:
        logging.error(f"Error converting {input_path} to .txt: {e}")
        raise

def process_folder(folder_path):
    """Convert PDF, DOC, DOCX to TXT with error handling and keep track of mapping."""
    processed_files = set()
    filename_mapping = {}
    for filename in os.listdir(folder_path):
        if filename in processed_files:
            continue
        file_path = os.path.join(folder_path, filename)
        base_name = os.path.splitext(filename)[0]
        safe_base_name = re.sub(r'[^\w\s.-]', '_', base_name)
        safe_base_name = re.sub(r'\s+', '_', safe_base_name)
        txt_path = os.path.join(folder_path, f"{safe_base_name}.txt")
        processed_files.add(filename)
        filename_mapping[filename] = os.path.basename(txt_path)
        try:
            if filename.lower().endswith('.pdf'):
                convert_pdf_to_txt(file_path, txt_path)
                if os.path.exists(txt_path):
                    os.remove(file_path)
                else:
                    logging.warning(f"No text file created for {filename}")
            elif filename.lower().endswith('.doc'):
                docx_path = convert_doc_to_docx(file_path, folder_path)
                if docx_path.endswith('.docx'):
                    text = extract_text_from_attachment(open(docx_path, 'rb').read(), docx_path)
                    if text:
                        with open(txt_path, 'w', encoding='utf-8') as f:
                            f.write(text)
                        logging.info(f"Converted {docx_path} to {txt_path}")
                    if docx_path != file_path:
                        os.remove(docx_path)
                os.remove(file_path)
                if not os.path.exists(txt_path):
                    logging.warning(f"No text file created for {filename}")
            elif filename.lower().endswith('.docx'):
                convert_docx_to_txt(file_path, txt_path)
                if os.path.exists(txt_path):
                    os.remove(file_path)
                    logging.info(f"Converted {file_path} to {txt_path}")
                else:
                    logging.warning(f"No text file created for {filename}")
        except Exception as e:
            logging.error(f"Error processing {filename}: {e}")
            continue
    txt_files = [f for f in os.listdir(folder_path) if f.endswith('.txt')]
    logging.info(f"Processed folder, found {len(txt_files)} .txt files")
    logging.info(f"Filename mapping: {filename_mapping}")
    return filename_mapping

def extract_year_of_birth_and_visa_from_email_body(email_body_text):
    """Extract Year of Birth and Visa Status from candidate's email body text."""
    yob = "N/A"
    visa = "N/A"

    # Enhanced patterns to match the format in your screenshot
    yob_pattern = r"(?:Year of Birth|DOB|Date of Birth|YOB)[:\s]*\(?19\d{2}\)?"
    visa_pattern = r"(?:Visa type|Visa Status|Visa)[:\s]*(USC|H1B|GC|OPT|L1|TN|EAD|Citizen|Permanent Resident)"

    yob_match = re.search(yob_pattern, email_body_text, re.IGNORECASE)
    if yob_match:
        # Extract just the 4-digit year
        year_match = re.search(r'19\d{2}', yob_match.group(0))
        if year_match:
            yob = year_match.group(0)

    visa_match = re.search(visa_pattern, email_body_text, re.IGNORECASE)
    if visa_match:
        visa = visa_match.group(1).strip()

    return yob, visa

def extract_candidate_details_from_resume_text(resume_text):
    """Extract candidate details from resume text using Groq LLM."""
    # Truncate resume text to avoid hitting token limits
    max_length = 10000  # Conservative limit for Llama3-8b
    if len(resume_text) > max_length:
        resume_text = resume_text[:max_length] + "... [truncated]"
    
    prompt = f"""
    Analyze the following resume text and extract the requested details. 
    Return ONLY a JSON object with the following structure:
    
    {{
        "name": "Full Name (combine first, middle if available, and last names)",
        "current_location": "Current location (city, state or country)",
        "total_experience": float (total years of experience as a decimal),
        "certification_count": integer (count of certifications),
        "government_work": {{
            "worked_with_govt": boolean,
            "govt_entities": ["list of government entities if any"]
        }},
        "skills": ["list of technical skills mentioned"]
    }}
    Special Instructions for SAP Skills:
    1. For each JD skill, count how many times it appears in the resume
    2. Include partial matches (e.g., 'Fiori' matches 'SAP Fiori')
    3. Group similar skills together
    # In your prompt to the LLM:
    "Special Instructions for Government Work Detection:\n"
    "1. ONLY identify government entities that start with 'Department of' or 'State of'\n"
   "2. Ignore all other government-related terms like Federal, City of, County of, etc.\n"
   "3. Mark worked_with_govt as true only if 'Department of' or 'State of' entities are found\n"
    Rules:
    1. Calculate total_experience by summing all work experience durations
    2. Count all certifications mentioned in education/certifications sections
    3. For government work, mark true if any government entity is mentioned
    4. Include only technical/professional skills, not soft skills
    
    Resume Text:
    {resume_text}
    """
    
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=GROQ_MODEL,
            temperature=0,
            response_format={"type": "json_object"},
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content
        response_text = re.sub(r'^```json\s*|\s*```$', '', response_text, flags=re.MULTILINE)
        
        try:
            details = json.loads(response_text)
        except json.JSONDecodeError:
            # Try to fix common JSON issues
            response_text = re.sub(r'(?<!\\)"(?!\s*[:}\],])', r'\"', response_text)
            response_text = re.sub(r"'(?!\s*[:}\],])", r"'", response_text)
            details = json.loads(response_text)
        
        # Format government work experience
        govt_work = "No"
        if details.get("government_work", {}).get("worked_with_govt", False):
            govt_entities = details["government_work"].get("govt_entities", [])
            if govt_entities:
                govt_work = "Yes: " + ", ".join(govt_entities)
        
        # Format skills list
        skills = details.get("skills", [])
        if isinstance(skills, str):
            skills = [s.strip() for s in skills.split(',') if s.strip()]
        
        return {
            "Name": details.get("name", "N/A").strip(),
            "Current Location": details.get("current_location", "N/A").strip(),
            "Experience": f"{float(details.get('total_experience', 0)):.2f} years",
            "Certification Count": int(details.get("certification_count", 0)),
            "Government Work": govt_work,
            "Skills": skills[:20]  # Limit to top 20 skills
        }
    except Exception as e:
        logging.error(f"Error extracting details from resume text: {e}")
        return {
            "Name": "N/A",
            "Current Location": "N/A",
            "Experience": "0.00 years",
            "Certification Count": 0,
            "Government Work": "No",
            "Skills": []
        }

def extract_resume_details(folder_path):
    """
    Extract candidate details from TXT resumes using Groq.
    """
    results = []
    
    for filename in os.listdir(folder_path):
        if not filename.endswith('.txt'):
            continue
            
        file_path = os.path.join(folder_path, filename)
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                resume_text = file.read()

            # Extract candidate details from resume text
            candidate_details = extract_candidate_details_from_resume_text(resume_text)
            
            results.append({
                'Filename': filename,
                'Name': candidate_details['Name'],
                'Current Location': candidate_details['Current Location'],
                'Experience': candidate_details['Experience'],
                'Certification Count': candidate_details['Certification Count'],
                'Government Work': candidate_details['Government Work'],
                'Skills': candidate_details['Skills']
            })
        except Exception as e:
            logging.error(f"Error processing {filename}: {e}")
            results.append({
                'Filename': filename,
                'Name': "N/A",
                'Current Location': "N/A",
                'Experience': "0.00 years",
                'Certification Count': 0,
                'Government Work': "No",
                'Skills': []
            })

    return pd.DataFrame(results)

def extract_email_data(service, message_id):
    """
    Extract candidate year of birth and visa status from email body.
    Extract candidate resume filename from attachments.
    """
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg.get("payload", {})

        # Extract email body text for candidate data
        email_body_text = extract_email_body(payload)

        # Extract Year of Birth and Visa Status from email body
        year_of_birth, visa_status = extract_year_of_birth_and_visa_from_email_body(email_body_text)

        # Extract attachments and identify resume filename
        attachments = fetch_attachments(service, message_id)
        resume_filename = identify_resume(service, message_id, attachments)

        details = {
            "Resume File": resume_filename,
            "Year of Birth": year_of_birth,
            "Visa Status": visa_status
        }

        return details

    except Exception as e:
        logging.error(f"Error processing email {message_id}: {e}")
        return {}

def apply_resume_scoring(df):
    """
    Apply ranking with priority: Government > Skills > Experience
    """
    logging.info(f"Applying resume scoring. DataFrame columns: {list(df.columns)}")
    
    # Initialize required columns if they don't exist
    if 'Composite Score' not in df.columns:
        df['Composite Score'] = 0.0
    if 'Experience Years' not in df.columns:
        df['Experience Years'] = 0.0
    if 'Government Score' not in df.columns:
        df['Government Score'] = 0
    if 'Matching Skills Count' not in df.columns:
        df['Matching Skills Count'] = 0
    if 'Rank' not in df.columns:
        df['Rank'] = 0
    
    def parse_experience(experience_str):
        if pd.isna(experience_str) or not isinstance(experience_str, str) or experience_str == "0 years":
            return 0.0
        match = re.match(r"(\d+\.?\d*)\s*years?", experience_str)
        if match:
            return float(match.group(1))
        return 0.0

    def calculate_government_score(government_work):
        if pd.isna(government_work) or government_work == "No":
            return 0
        if "Yes:" in government_work:
            entities = government_work.split("Yes:")[1].split(",")
            return len(entities) * 1000  # Heavy weighting for government
        return 1000  # Heavy weighting for government work without specific entities

    try:
        # Parse values
        df["Experience Years"] = df["Experience"].apply(parse_experience)
        df["Government Score"] = df["Government Work"].apply(calculate_government_score)
        
        # Get max values for normalization (avoid division by zero)
        max_exp = df["Experience Years"].max() or 1
        max_skills = df["Matching Skills Count"].max() or 1
        
        # NEW PRIORITY-BASED SCORING: Government > Skills > Experience
        # Government gets massive weight (1000+), Skills gets moderate weight (100+), Experience gets base weight
        df["Composite Score"] = (
            df["Government Score"] +  # Primary: Government (heaviest weight)
            (df["Matching Skills Count"] / max_skills * 100) +  # Secondary: Skills
            (df["Experience Years"] / max_exp)   # Tertiary: Experience
        )
        
        # Assign ranks based on composite score (higher score = better rank)
        df["Rank"] = df["Composite Score"].rank(ascending=False, method="min").astype(int)
        
        # Drop temporary columns
        df = df.drop(columns=["Composite Score", "Experience Years", "Government Score"], errors="ignore")
        
        # Sort by rank
        df = df.sort_values(by="Rank", ascending=True).reset_index(drop=True)
        
        return df, "Scenario: Government (Primary) > Skills (Secondary) > Experience (Tertiary)"
    
    except Exception as e:
        logging.error(f"Error in resume scoring: {e}")
        # Fallback to simple ranking if scoring fails
        df["Rank"] = range(1, len(df)+1)
        return df, "Scenario: Simple Ranking (Fallback)"

def extract_skills_from_subject(subject):
    """
    Extracts skills dynamically from the email subject, removing location-related words (Hybrid, Local, Remote).
    Splits slash-separated multi-skills like AIX/Unix/Linux, but keeps compound terms like Z/OS or CI/CD intact.
    Returns the skills as a Python list.
    """
    # Remove 'Fwd:' specifically from the beginning of the subject
    subject = re.sub(r'^\s*Fwd:\s*', '', subject, flags=re.IGNORECASE)
    # Remove location-related keywords
    subject = re.sub(r'\b(Hybrid|Local|Remote|Onsite)\b[\s/]*', '', subject, flags=re.IGNORECASE).strip()
    # Remove any extra spaces or leftover punctuation
    subject = re.sub(r'\s{2,}', ' ', subject)
    subject = re.sub(r'^[,\s]+|[,\s]+$', '', subject)
 
    # Split at 'with' (case insensitive)
    split_parts = re.split(r'\bwith\b', subject, flags=re.IGNORECASE)
 
    if len(split_parts) == 2:
        before_with = split_parts[0].strip()
        after_with = split_parts[1].strip()
 
        parts = [part.strip() for part in after_with.split(',') if part.strip()]
        skills = []
 
        for part in parts:
            # If it's a compound like Z/OS, CI/CD, keep it as-is
            if re.match(r'^[A-Za-z0-9]+/[A-Za-z0-9]+$', part):
                skills.append(part)
            else:
                # Split AIX/Unix/Linux into ['AIX', 'Unix', 'Linux']
                skills.extend([s.strip() for s in part.split('/') if s.strip()])

        return [before_with] + skills
    else:
        # No 'with' found, fallback to comma-splitting
        skills_raw = [s.strip() for s in subject.split(',') if s.strip()]
        return skills_raw

def extract_subject_skills(text, use_subject=True):
    """
    Extracts the skills list from an email body. First tries the subject line, then falls back to content.
    Returns a comma-separated string of skills or None if nothing is found.
    """
    if use_subject:
        # First try to extract from the subject line
        subject_match = re.search(r'^Subject:\s*(.+)$', text, re.IGNORECASE | re.MULTILINE)
        if subject_match:
            subject = subject_match.group(1)
            subject_skills = extract_skills_from_subject(subject)
            if subject_skills:
                return ', '.join(subject_skills)
 
    # Fallback: scan the email body
    lines = [line.strip() for line in text.split('\n') if line.strip()]
 
    patterns = [
        r'Job ID:',                          # Job ID: TX-70125059
        r'# \d{1,2} \w+, \d{4}',             # # 11 April, 2025
        r'\b\d{1,2}/\d{1,2}/\d{4}\b',        # 4/11/2025
        r'Posted on \w+ \d{1,2}, \d{4}'      # Posted on April 11, 2025
    ]
 
    for i, line in enumerate(lines):
        if any(re.search(pattern, line, re.IGNORECASE) for pattern in patterns):
            # Look ahead to next 3 lines
            for j in range(i + 1, min(i + 4, len(lines))):
                next_line = lines[j]
                if (next_line and
                    not re.search(r'^(By|Location|Remote|Hybrid|Onsite|Read on)', next_line, re.IGNORECASE) and
                    len(next_line) > 20):
                    return next_line

    return None

def extract_job_description_skills(email_body):
    """
    Extract skills from job description in email body.
    """
    # First try to find a clear skills section
    skills_section = extract_skills(email_body)
    if skills_section:
        return ', '.join(skills_section)
    
    # Fallback to looking for bulleted lists that might be skills
    bullet_items = re.findall(r'(?:\n\s*[-•*]\s*)(.+?)(?=\n\s*[-•*]|$)', email_body)
    if bullet_items:
        return ', '.join(bullet_items)
    
    return None

def identify_job_description_email(service, messages):
    """
    Identify the job description email from a list of messages.
    Returns the message ID of the job description email or None if not found.
    """
    for msg in messages:
        message_id = msg["id"]
        try:
            msg_data = service.users().messages().get(userId="me", id=message_id, format="full").execute()
            payload = msg_data.get("payload", {})
            headers = payload.get("headers", [])
            subject = next((h['value'] for h in headers if h['name'] == 'Subject'), "").lower()
            
            # Skip candidate emails (usually have "resume" or "candidate" in subject)
            if "resume" in subject or "candidate" in subject:
                continue
                
            # Check if this looks like a job description email
            email_body = extract_email_body(payload)
            if not email_body:
                continue
                
            # Look for indicators of a job description
            jd_indicators = [
                "job description", "job posting", "job id", "position overview",
                "required skills", "required experience", "skills required",
                "qualifications", "responsibilities", "job requirements"
            ]
            
            if any(indicator in email_body.lower() for indicator in jd_indicators):
                return message_id
                
        except Exception as e:
            logging.error(f"Error checking message {message_id}: {e}")
            continue
            
    return None

def extract_job_description_details(service, message_id):
    """
    Extract detailed job description information from the email.
    Returns a dictionary with job role, subject skills, and JD skills.
    """
    if not message_id:
        return None
        
    try:
        msg_data = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg_data.get("payload", {})
        headers = payload.get("headers", [])
        subject = next((h['value'] for h in headers if h['name'] == 'Subject'), "")
        email_body = extract_email_body(payload)
        
        # Extract job role from subject (before "with" if present)
        job_role = subject.split("With")[0].split("with")[0].strip()
        
        # Clean up job role (remove location info)
        job_role = re.sub(r'\b(Hybrid|Local|Remote|Onsite)\b[\s/]*', '', job_role, flags=re.IGNORECASE).strip()
        job_role = re.sub(r'\(\d+\+\)', '', job_role).strip()  # Remove experience like (12+)
        
        # Extract skills from subject
        subject_skills = extract_skills_from_subject(subject)
        
        # Extract skills from job description body
        jd_skills = extract_job_description_skills(email_body)
        
        # Enhanced JD skills extraction
        if not jd_skills or jd_skills == "N/A":
            # Try to find a skills section with years of experience
            skills_section = extract_skills_with_experience(email_body)
            if skills_section:
                jd_skills = skills_section
        
        return {
            "Job Role": job_role,
            "Subject Skills": subject_skills,
            "JD Skills": jd_skills,
            "Full Job Description": email_body  # Keep full text for reference
        }
        
    except Exception as e:
        logging.error(f"Error extracting job description from {message_id}: {e}")
        return None

def extract_skills_with_experience(email_body):
    """
    Extract skills section that includes years of experience requirements.
    Returns formatted skills text or None if not found.
    """
    # Look for a clear skills section with experience years
    skill_patterns = [
        r"Skills?:\s*(.*?)(?:\n\s*[A-Z][a-z]+:|$)",  # Skills: ... next section
        r"Skills?\s*&\s*Qualifications?:\s*(.*?)(?:\n\s*[A-Z][a-z]+:|$)",
        r"Technical\s*Skills?:\s*(.*?)(?:\n\s*[A-Z][a-z]+:|$)",
        r"Required\s*Skills?:\s*(.*?)(?:\n\s*[A-Z][a-z]+:|$)"
    ]
    
    for pattern in skill_patterns:
        match = re.search(pattern, email_body, re.DOTALL | re.IGNORECASE)
        if match:
            skill_text = match.group(1).strip()
            # Clean up the skills text
            skill_text = re.sub(r'\s+', ' ', skill_text)  # Normalize whitespace
            skill_text = re.sub(r'(\n\s*[-•*]\s*)', '\n', skill_text)  # Remove bullets
            return skill_text
            
    return None

def count_matching_skills(candidate_skills, jd_skills):
    """
    Enhanced skill matching with fuzzy matching and SAP-specific handling
    Returns dictionary of {skill: count} and total matches.
    """
    if not candidate_skills or not jd_skills:
        return {}, 0
    
    # Normalize candidate skills
    candidate_skills = [str(skill).lower().strip() for skill in candidate_skills if pd.notna(skill)]
    
    # Process JD skills with special handling for SAP terms
    if isinstance(jd_skills, str):
        # First split by commas but preserve SAP compound terms
        jd_skills = re.split(r',(?![^/]*/)', jd_skills)
        jd_skills = [skill.strip() for skill in jd_skills if skill.strip()]
        
        # List of SAP terms that shouldn't be split
        sap_compound_terms = [
            'sap fiori', 'ui5', 'odata', 'cds views', 
            'gateway service', 'web dynpro', 'hana'
        ]
        
        final_jd_skills = []
        for skill in jd_skills:
            skill_lower = skill.lower()
            # Keep compound terms together
            if any(term in skill_lower for term in sap_compound_terms):
                final_jd_skills.append(skill_lower)
            else:
                # Split simple slash-separated skills
                final_jd_skills.extend([s.strip() for s in skill_lower.split('/') if s.strip()])
        jd_skills = final_jd_skills
    elif isinstance(jd_skills, list):
        jd_skills = [str(skill).lower().strip() for skill in jd_skills if pd.notna(skill)]
    else:
        jd_skills = []
    
    # Count matches with improved logic
    skill_counts = {}
    total_matches = 0
    
    for c_skill in candidate_skills:
        count = 0
        for jd_skill in jd_skills:
            # Exact match
            if c_skill == jd_skill:
                count += 1
            # Partial match (e.g., "fiori" in "sap fiori")
            elif c_skill in jd_skill or jd_skill in c_skill:
                count += 1
            # Fuzzy match for similar terms
            elif fuzz.ratio(c_skill, jd_skill) > 80:
                count += 1
        
        if count > 0:
            skill_counts[c_skill] = count
            total_matches += count
    
    return skill_counts, total_matches

def filter_and_count_skills(df):
    """Add debug logging for skill matching with proper error handling"""
    logging.info("Starting skill matching process...")
    
    # Initialize columns if they don't exist
    if 'Matching Skills' not in df.columns:
        df['Matching Skills'] = [[] for _ in range(len(df))]
    if 'Matching Skills Count' not in df.columns:
        df['Matching Skills Count'] = 0
    
    for idx, row in df.iterrows():
        try:
            logging.info(f"\nProcessing candidate: {row['Name']}")
            logging.info(f"Candidate skills: {row.get('Skills', [])}")
            
            jd_skills = []
            if isinstance(row.get('Subject Skills'), (str, list)) and row['Subject Skills'] != "N/A":
                logging.info(f"Subject Skills: {row['Subject Skills']}")
                jd_skills.append(row['Subject Skills'])
            if isinstance(row.get('JD Skills'), (str, list)) and row['JD Skills'] != "N/A":
                logging.info(f"JD Skills: {row['JD Skills']}")
                jd_skills.append(row['JD Skills'])
            
            if isinstance(row.get('Skills'), list):
                candidate_skills = [s for s in row['Skills'] if pd.notna(s)]
                logging.info(f"Raw candidate skills: {candidate_skills}")
                
                matching_skills, total_matches = count_matching_skills(candidate_skills, jd_skills)
                logging.info(f"Matched skills: {matching_skills}")
                
                formatted_skills = [f"{skill} ({count})" for skill, count in matching_skills.items()]
                
                df.at[idx, 'Matching Skills'] = formatted_skills
                df.at[idx, 'Matching Skills Count'] = total_matches
            else:
                df.at[idx, 'Matching Skills'] = []
                df.at[idx, 'Matching Skills Count'] = 0
                
        except Exception as e:
            logging.error(f"Error processing skills for candidate {row.get('Name', 'Unknown')}: {e}")
            df.at[idx, 'Matching Skills'] = []
            df.at[idx, 'Matching Skills Count'] = 0
    
    return df

def store_results_in_supabase(supabase_client, job_id, df, jd_details):
    """Store resume analysis results in Supabase with proper data structure"""
    try:
        # Clean the DataFrame first
        df = df.replace([np.nan, 'nan', 'N/A', 'Unknown'], None)
        
        # Convert numeric columns, filling NaN with 0
        numeric_cols = ['Rank', 'Certification Count', 'Matching Skills Count']
        for col in numeric_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).astype(int)
        
        # Prepare data for insertion
        records = df.to_dict('records')
        data_to_insert = []
        current_time = datetime.now(pytz.UTC).isoformat()
        
        for record in records:
            # Safely handle all fields with proper defaults
            prepared_record = {
                "job_id": job_id,
                "rank": int(record.get("Rank", 0)) if pd.notna(record.get("Rank")) else 0,
                "name": str(record.get("Name", "")).strip() or "Unknown",
                "current_location": str(record.get("Current Location", "")).strip() or "Unknown",
                "government_work": str(record.get("Government Work", "")).strip() or "No",
                "experience": str(record.get("Experience", "")).strip() or "0 years",
                "certification_count": int(record.get("Certification Count", 0)) if pd.notna(record.get("Certification Count")) else 0,
                "matching_skills": ', '.join(record.get("Matching Skills", [])) if isinstance(record.get("Matching Skills"), list) else str(record.get("Matching Skills", "")),
                "matching_skills_count": int(record.get("Matching Skills Count", 0)) if pd.notna(record.get("Matching Skills Count")) else 0,
                "job_role": str(jd_details.get("Job Role", "")).strip() or "N/A",
                "subject_skills": jd_details.get("Subject Skills", []) or [],
                "jd_skills": str(jd_details.get("JD Skills", "")).strip() or "N/A",
                "created_at": current_time
            }
            
            # Clean empty strings and lists
            for key, value in prepared_record.items():
                if isinstance(value, str) and not value.strip():
                    prepared_record[key] = "N/A"
                elif isinstance(value, list) and not value:
                    prepared_record[key] = []
            
            data_to_insert.append(prepared_record)

        # Insert in batches with error handling
        batch_size = 50
        total_inserted = 0
        errors = []
        
        for i in range(0, len(data_to_insert), batch_size):
            batch = data_to_insert[i:i + batch_size]
            try:
                insert_response = supabase_client.table("resume_analysis") \
                    .insert(batch) \
                    .execute()
                
                if hasattr(insert_response, 'data'):
                    total_inserted += len(insert_response.data)
                else:
                    errors.append(f"Batch {i//batch_size}: No data in response")
            except Exception as e:
                errors.append(f"Batch {i//batch_size}: {str(e)}")
                logging.error(f"Error inserting batch {i//batch_size}: {e}")
                # Try inserting records one by one
                for record in batch:
                    try:
                        supabase_client.table("resume_analysis").insert(record).execute()
                        total_inserted += 1
                    except Exception as single_error:
                        errors.append(f"Single record failed: {str(single_error)}")
                        logging.error(f"Failed to insert single record: {single_error}")

        if errors:
            logging.warning(f"Completed with {len(errors)} errors")
            return True, f"Stored {total_inserted} records with {len(errors)} errors"
        
        logging.info(f"Successfully stored {total_inserted} records for job {job_id}")
        return True, f"Stored {total_inserted} records for job {job_id}"
        
    except Exception as e:
        error_msg = f"Failed to store results in Supabase: {str(e)}"
        logging.error(error_msg)
        # Try to save at least the CSV for debugging
        try:
            df.to_csv(f"failed_upload_{job_id}.csv", index=False)
            logging.info(f"Saved failed data to failed_upload_{job_id}.csv")
        except:
            pass
        return False, error_msg

def get_supabase_client():
    """Initialize and verify Supabase connection with retry logic"""
    max_retries = 3
    retry_delay = 1  # seconds
    
    for attempt in range(max_retries):
        try:
            client = create_client(SUPABASE_URL, SUPABASE_KEY)
            
            # Test connection with a simple query (removed timeout parameter)
            response = client.table("resume_analysis") \
                .select("job_id", count='exact') \
                .limit(1) \
                .execute()
            
            if response is None:
                raise ConnectionError("No response from Supabase")
                
            logging.info("Supabase connection verified")
            return client
            
        except Exception as e:
            if attempt == max_retries - 1:
                logging.error(f"Supabase connection failed after {max_retries} attempts: {e}")
                raise ConnectionError(f"Could not connect to Supabase after {max_retries} attempts") from e
                
            logging.warning(f"Supabase connection attempt {attempt + 1} failed, retrying...")
            time.sleep(retry_delay * (attempt + 1))

def remove_duplicate_candidates(df):
    """Remove duplicate candidates based on name, location, and experience."""
    # Create a hash for each candidate based on key fields
    df['CandidateHash'] = df.apply(lambda row: hashlib.md5(
        f"{row['Name']}{row['Current Location']}{row['Experience']}{row['Year of Birth']}".encode()
    ).hexdigest(), axis=1)
    
    # Keep the first occurrence of each duplicate
    df = df.drop_duplicates(subset=['CandidateHash'], keep='first')
    
    # Recalculate ranks after removing duplicates
    if 'Composite Score' in df.columns:
        df['Rank'] = df['Composite Score'].rank(ascending=False, method='min').astype(int)
    else:
        df['Rank'] = range(1, len(df)+1)
    
    return df.drop(columns=['CandidateHash'])

def main(job_id):
    try:
        service = auto_authenticate_google()
        logging.info(f"Starting processing for job ID: {job_id}")
        resume_folder = create_resume_folder(RESUME_FOLDER)
        messages = get_emails_by_job_id(service, job_id)
        
        if not messages:
            print(f"No emails found related to job ID {job_id}.")
            logging.info(f"No emails found for job ID {job_id}. Terminating.")
            return

        print(f"Found {len(messages)} emails related to job ID {job_id}:")
        logging.info(f"Processing {len(messages)} emails")
        
        # First identify the job description email
        jd_message_id = identify_job_description_email(service, messages)
        jd_details = extract_job_description_details(service, jd_message_id) if jd_message_id else None
        
        if jd_details:
            print("\nJob Description Details:")
            print(f"Job Role: {jd_details['Job Role']}")
            print(f"Subject Skills: {', '.join(jd_details['Subject Skills']) if jd_details['Subject Skills'] else 'N/A'}")
            print("\nJD Skills:")
            print(jd_details['JD Skills'] if jd_details['JD Skills'] else 'N/A')
        else:
            print("\nNo job description email found for this job ID.")
            jd_details = {
                "Job Role": "N/A",
                "Subject Skills": [],
                "JD Skills": "N/A",
                "Full Job Description": "N/A"
            }
        
        # Process candidate emails
        email_data = []
        failure_reasons = []

        for msg in messages:
            message_id = msg["id"]
            
            if message_id == jd_message_id:
                continue
                
            try:
                msg_data = service.users().messages().get(userId="me", id=message_id, format="full").execute()
                payload = msg_data.get("payload", {})
                headers = payload.get("headers", [])
                subject = next((h['value'] for h in headers if h['name'] == 'Subject'), "")
                email_body_text = extract_email_body(payload)
                details = extract_email_data(service, message_id)
                
                if details:
                    details.update({
                        "Job Role": jd_details["Job Role"],
                        "Subject Skills": ', '.join(jd_details["Subject Skills"]) if jd_details["Subject Skills"] else "N/A",
                        "JD Skills": jd_details["JD Skills"],
                        "Full Job Description": jd_details["Full Job Description"],
                        "Candidate Email Subject": subject
                    })
                    email_data.append(details)
                    
                    attachments = fetch_attachments(service, message_id)
                    if attachments:
                        save_resumes_to_folder(service, details, message_id, attachments, resume_folder)
                    else:
                        logging.info(f"No attachments found for message {message_id}")
                        failure_reasons.append(f"Message {message_id}: No attachments found")
                else:
                    logging.warning(f"No candidate details extracted for message {message_id}")
                    failure_reasons.append(f"Message {message_id}: Failed to extract candidate details")
                    
            except Exception as e:
                logging.error(f"Error processing message {message_id}: {e}")
                failure_reasons.append(f"Message {message_id}: Processing error - {str(e)}")
                continue

        logging.info(f"Total candidates extracted: {len(email_data)}")

        if email_data:
            df = pd.DataFrame(email_data)
            filename_mapping = process_folder(RESUME_FOLDER)
            resume_df = extract_resume_details(RESUME_FOLDER)
            
            def normalize_filename(filename):
                if pd.isna(filename) or filename == "N/A":
                    return "N/A"
                base = os.path.splitext(filename)[0]
                base = re.sub(r'[^\w\s.-]', '_', base)
                base = re.sub(r'\s+', '_', base)
                return f"{base}.txt"

            df["Resume File"] = df["Resume File"].apply(normalize_filename)
            resume_df["Filename"] = resume_df["Filename"].apply(lambda x: x if pd.notna(x) else "N/A")
            
            df = pd.merge(df, resume_df, left_on="Resume File", right_on="Filename", how="left")
            df = df[~((df['Name'] == 'N/A') & (df['Current Location'] == 'N/A') & (df['Experience'] == '0.00 years'))]
            
            df, scenario = apply_resume_scoring(df)
            
            # Remove duplicate candidates before final processing
            df = remove_duplicate_candidates(df)
            
            print(f"\nRanking Criteria: {scenario}")
            
            df = filter_and_count_skills(df)
            df = df.sort_values(by=['Rank', 'Matching Skills Count'], ascending=[True, False])
            
            output_columns = [
                "Rank", "Name", "Current Location", "Year of Birth", "Visa Status",
                "Experience", "Certification Count", "Government Work", "Job Role",
                "Subject Skills", "JD Skills", "Matching Skills", "Matching Skills Count",
                "Resume File", "Candidate Email Subject"
            ]
            available_columns = [col for col in output_columns if col in df.columns]
            df_output = df[available_columns]
            
            for col in ["Matching Skills", "Skills"]:
                if col in df_output.columns:
                    df_output[col] = df_output[col].apply(lambda x: ', '.join(x) if isinstance(x, list) else x)
            
            # Save to CSV
            df_output.to_csv(OUTPUT_CSV, index=False)
            logging.info(f"Results saved to {OUTPUT_CSV}")

            # Store in Supabase
            try:
                sb = get_supabase_client()
                store_results_in_supabase(sb, job_id, df_output, jd_details)
            except Exception as e:
                logging.error(f"Storage failed: {e}")

            # Display results
            display_columns = [
                "Rank", "Name", "Current Location", "Government Work",
                "Experience", "Certification Count", "Matching Skills", "Matching Skills Count"
            ]
            display_columns = [col for col in display_columns if col in df.columns]
            
            print("\nCandidate Summary (Sorted by Rank):")
            print(tabulate(df[display_columns], headers="keys", tablefmt="grid", showindex=False))
            
    except Exception as e:
        print(f"Failed to process resumes: {e}")
        logging.error(f"Failed to process resumes: {e}")
        return

if __name__ == "__main__":
    job_id = input("Enter the job ID to search for: ").strip()
    main(job_id)